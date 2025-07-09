import streamlit as st
import pandas as pd
import numpy as np
import datetime
import requests
import os
from datetime import datetime, time
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
import io
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Import Google Drive utility functions
from drive_utils import list_csv_files_in_folder, download_csv_to_df

# DeepSeek API configuration
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "sk-8fa69e0815904d0daf8831374c9999b8")
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"
MODEL = "deepseek-chat"

# --- Sidebar settings ---
st.sidebar.title("Settings")

# REQUEST 1: Renamed the input and it will be used to define "Low Speed Events"
mc_min_speed_req = st.sidebar.number_input("M/c. Min Speed Req.", value=10)

idle_time = st.sidebar.number_input("Idle Time Threshold (min)", value=10)
user_machine_map = st.sidebar.text_area(
    "User-Machine Mapping (user:machine1,machine2)",
    value="user1:Extruder-80A\nuser2:Extruder-90B"
)

# Add a refresh button to the sidebar
if st.sidebar.button("Refresh File List"):
    st.cache_data.clear() # Clear the cache to force refetching file list
    st.rerun()

# --- DATA LOADING FROM GOOGLE DRIVE ---
st.title("IoT Machine Analytics Dashboard (Google Drive & DeepSeek AI)")

# FIXED: Hardcoded the Google Drive Folder ID to resolve the environment variable error.
GOOGLE_DRIVE_FOLDER_ID = "1r75YeqkkeskfjXRPcf9salnBb4DrhvGV"

if not GOOGLE_DRIVE_FOLDER_ID:
    st.error("Fatal Error: GOOGLE_DRIVE_FOLDER_ID is not set.")
    st.stop()

# Get list of CSV files from the specified Google Drive folder
try:
    with st.spinner("Connecting to Google Drive and fetching file list..."):
        drive_files = list_csv_files_in_folder(GOOGLE_DRIVE_FOLDER_ID)
except Exception as e:
    st.error(f"Could not connect to Google Drive. Please check your credentials and Folder ID.")
    st.error(f"Details: {e}")
    st.stop()

if not drive_files:
    st.warning("No CSV files found in the specified Google Drive folder. Please add files to continue.")
    st.stop()

# Display the number of files uploaded
st.info(f"Number of files found in Google Drive: **{len(drive_files)}**")

# Create a mapping of file names to their unique Google Drive IDs
file_map = {file['name']: file['id'] for file in drive_files}
file_names = list(file_map.keys())

# Sidebar selection for the file to be analyzed
selected_file = st.sidebar.selectbox("Select a file for analysis", file_names)

# Load the selected file from Google Drive into a DataFrame
if selected_file:
    try:
        with st.spinner(f"Downloading and loading '{selected_file}'..."):
            file_id = file_map[selected_file]
            df = download_csv_to_df(file_id)
            df.columns = [c.strip().lower() for c in df.columns]
            df['timestamp'] = pd.to_datetime(df['timestamp'], errors='coerce')
            df = df.dropna(subset=['timestamp']).sort_values('timestamp')
    except Exception as e:
        st.error(f"Failed to download or process the file '{selected_file}'.")
        st.error(f"Details: {e}")
        st.stop()
else:
    st.info("Please select a file from the sidebar to begin analysis.")
    st.stop()

# --- Core Definitions based on new logic ---

# REQUEST 2 & 3: Idle is now strictly speed <= 1. This also ensures no speed > 1 is in an idle event.
df['idle'] = df['speed'] <= 1

# Running is now defined as any speed greater than 1.
df['is_running'] = df['speed'] > 1

# REQUEST 1: A "low speed" state is when speed is below the requirement, but not idle (i.e., > 1).
df['is_low_speed'] = (df['speed'] < mc_min_speed_req) & (df['speed'] > 1)

# Add time-based columns for analysis
df['date'] = df['timestamp'].dt.date
df['time'] = df['timestamp'].dt.time
df['hour'] = df['timestamp'].dt.hour

# INITIALIZE VARIABLES EARLY
min_date = df['timestamp'].dt.date.min()
max_date = df['timestamp'].dt.date.max()

# KPI calculations (now based on the new definitions)
total_minutes = (df['timestamp'].iloc[-1] - df['timestamp'].iloc[0]).total_seconds() / 60 if len(df) > 1 else 0
running_minutes = df['is_running'].sum()
uptime_percent = 100 * running_minutes / total_minutes if total_minutes else 0
avg_speed = df.loc[df['is_running'], 'speed'].mean() if running_minutes else 0
total_production = df['quantity'].iloc[-1] - df['quantity'].iloc[0] if len(df) > 1 else 0

# --- DATA PROCESSING FOR TABLES ---

# Calculate Idle Events based on the new definition (speed <= 1)
df['idle_group'] = (df['idle'] != df['idle'].shift()).cumsum()
idle_events = df[df['idle']].groupby('idle_group').agg(
    start=('timestamp', 'first'),
    end=('timestamp', 'last'),
    duration=('timestamp', lambda x: (x.iloc[-1] - x.iloc[0]).total_seconds() / 60),
    avg_speed=('speed', 'mean')
)
long_idles = idle_events[idle_events['duration'] >= idle_time]

# Calculate Low Speed Events based on the new definition
df['low_speed_group'] = (df['is_low_speed'] != df['is_low_speed'].shift()).cumsum()
low_speed_events_all = df[df['is_low_speed']].groupby('low_speed_group').agg(
    start=('timestamp', 'first'),
    end=('timestamp', 'last'),
    duration=('timestamp', lambda x: round((x.iloc[-1] - x.iloc[0]).total_seconds() / 60, 4)),
    min_speed=('speed', 'min'),
    max_speed=('speed', 'max'),
    avg_speed=('speed', 'mean')
).reset_index(drop=True)

# Sort in descending order of duration
low_speed_events_sorted = low_speed_events_all.sort_values('duration', ascending=False)


# --- START OF UI DISPLAY ---

# Display KPIs at the top
st.subheader("Key Performance Indicators")
kpi1, kpi2, kpi3, kpi4 = st.columns(4)
kpi1.metric("Uptime (%)", f"{uptime_percent:.1f}")
kpi2.metric("Avg Speed (RPM)", f"{avg_speed:.1f}")
kpi3.metric("Production", f"{total_production:.1f}")
kpi4.metric(f"Idle Events (≥{idle_time} min)", f"{len(long_idles)}")

# Date/time filtering controls
st.subheader("Filter Data for Analysis")
col1, col2 = st.columns(2)
with col1:
    date_range = st.date_input("Select Date Range", value=(min_date, max_date), min_value=min_date, max_value=max_date, format="YYYY/MM/DD")
with col2:
    time_range = st.slider("Select Time Range (Hours)", min_value=0, max_value=23, value=(0, 23), step=1, format="%d:00")

# Filter data based on selections
if len(date_range) == 2:
    start_date, end_date = date_range
    filtered_df = df[(df['timestamp'].dt.date >= start_date) & (df['timestamp'].dt.date <= end_date) & (df['hour'] >= time_range[0]) & (df['hour'] <= time_range[1])].copy()
else:
    filtered_df = df[(df['timestamp'].dt.date == date_range[0]) & (df['hour'] >= time_range[0]) & (df['hour'] <= time_range[1])].copy()

st.info(f"Showing data from {filtered_df['timestamp'].min()} to {filtered_df['timestamp'].max()} ({len(filtered_df)} records)")


# --- TABLES SECTION ---
st.header("Data Tables")

# Low Speed Events Table
st.subheader("Low Speed Events")
st.markdown(f"Events where speed was between **1 RPM** and the **{mc_min_speed_req} RPM** requirement. Sorted by duration. Based on the **full dataset**.")
st.dataframe(low_speed_events_sorted)

# Idle Events Analysis Table
st.subheader("Idle Events Analysis")
st.markdown(f"Events where the machine speed was **≤ 1 RPM** for longer than the **{idle_time} min threshold**. Based on the **filtered date/time range** above.")
display_idle_events = idle_events[idle_events['duration'] >= idle_time]

if len(date_range) == 2:
    start_datetime = datetime.combine(date_range[0], time(time_range[0], 0))
    end_datetime = datetime.combine(date_range[1], time(time_range[1], 59, 59))
else:
    start_datetime = datetime.combine(date_range[0], time(time_range[0], 0))
    end_datetime = datetime.combine(date_range[0], time(time_range[1], 59, 59))

if not display_idle_events.empty:
    filtered_idle_events = display_idle_events[(display_idle_events['start'] >= start_datetime) & (display_idle_events['end'] <= end_datetime)].copy()
else:
    filtered_idle_events = pd.DataFrame()

if not filtered_idle_events.empty:
    filtered_idle_events = filtered_idle_events.sort_values('duration', ascending=False)
    filtered_idle_events['start_time'] = filtered_idle_events['start'].dt.strftime('%H:%M:%S')
    filtered_idle_events['end_time'] = filtered_idle_events['end'].dt.strftime('%H:%M:%S')
    filtered_idle_events['start_date'] = filtered_idle_events['start'].dt.date
    filtered_idle_events['idle_group'] = filtered_idle_events.index
    display_columns = ['idle_group', 'start_date', 'start_time', 'end_time', 'duration', 'avg_speed']
    st.dataframe(filtered_idle_events[display_columns])
else:
    st.info("No idle events found matching the criteria in the selected time range.")

# Hourly Performance Table
st.subheader("Hourly Performance Statistics")
st.markdown("Performance metrics aggregated by hour, based on the **filtered date/time range** above.")
hourly_stats = filtered_df.groupby('hour').agg({
    'speed': ['mean', 'max', 'min'],
    'is_running': 'mean',
    'quantity': lambda x: x.iloc[-1] - x.iloc[0] if len(x) > 1 else 0
}).round(2)
hourly_stats.columns = ['Avg_Speed', 'Max_Speed', 'Min_Speed', 'Uptime_Ratio', 'Production']
hourly_stats['Uptime_Percent'] = (hourly_stats['Uptime_Ratio'] * 100).round(1)
st.dataframe(hourly_stats[['Avg_Speed', 'Max_Speed', 'Min_Speed', 'Uptime_Percent', 'Production']])


# --- GRAPHS SECTION ---
st.header("Graphical Analysis")
col1, col2 = st.columns(2)
with col1:
    graph_view = st.selectbox("Select Graph View", ["Total Graph", "Running Only", "Idle Only", "Low Speed Events", "Production vs Speed", "Idle Events Timeline", "Hourly Performance Pattern"])
with col2:
    chart_type = st.selectbox("Select Chart Type", ["Bar Chart", "Line Chart", "Area Chart"])

# Graphing functions
def create_bar_chart(data, title, color='blue'):
    fig = go.Figure()
    fig.add_trace(go.Bar(x=data.index, y=data['speed'], name='Speed (RPM)', marker_color=color, text=data['speed'].round(1), textposition='outside'))
    fig.update_layout(title=title, xaxis_title="Time", yaxis_title="Speed (RPM)", showlegend=True, height=500, xaxis=dict(tickangle=45))
    return fig

def create_multi_bar_chart(data, columns, title):
    fig = go.Figure()
    colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd']
    for i, col in enumerate(columns):
        fig.add_trace(go.Bar(x=data.index, y=data[col], name=col, marker_color=colors[i % len(colors)]))
    fig.update_layout(title=title, xaxis_title="Time", yaxis_title="Value", barmode='group', showlegend=True, height=500, xaxis=dict(tickangle=45))
    return fig

# Graph rendering logic
if graph_view == "Total Graph":
    st.subheader("Speed Over Time - Complete View")
    if chart_type == "Bar Chart":
        chart_data = filtered_df.set_index('timestamp')[['speed']]
        fig = create_bar_chart(chart_data, "Machine Speed Over Time", '#1f77b4')
        st.plotly_chart(fig, use_container_width=True)
    elif chart_type == "Area Chart":
        st.area_chart(filtered_df.set_index('timestamp')[['speed']], use_container_width=True)
    else:
        st.line_chart(filtered_df.set_index('timestamp')[['speed']], use_container_width=True)
    
    st.subheader("Production Over Time")
    if chart_type == "Bar Chart":
        production_data = filtered_df.set_index('timestamp')[['quantity']]
        fig = go.Figure(go.Bar(x=production_data.index, y=production_data['quantity'], name='Production', marker_color='#2ca02c'))
        fig.update_layout(title="Production Quantity Over Time", xaxis_title="Time", yaxis_title="Quantity")
        st.plotly_chart(fig, use_container_width=True)
    elif chart_type == "Area Chart":
        st.area_chart(filtered_df.set_index('timestamp')[['quantity']], use_container_width=True)
    else:
        st.line_chart(filtered_df.set_index('timestamp')[['quantity']], use_container_width=True)

elif graph_view == "Running Only":
    st.subheader("Speed Over Time - Running Periods Only (Speed > 1 RPM)")
    running_data = filtered_df[filtered_df['is_running']].copy()
    if not running_data.empty:
        if chart_type == "Bar Chart":
            fig = create_bar_chart(running_data.set_index('timestamp')[['speed']], "Machine Speed - Running Periods", '#00ff00')
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.line_chart(running_data.set_index('timestamp')[['speed']], color='#00ff00', use_container_width=True)
    else:
        st.warning("No running periods found in the selected time range.")

elif graph_view == "Idle Only":
    st.subheader("Speed Over Time - Idle Periods Only (Speed <= 1 RPM)")
    idle_data = filtered_df[filtered_df['idle']].copy()
    if not idle_data.empty:
        if chart_type == "Bar Chart":
            fig = create_bar_chart(idle_data.set_index('timestamp')[['speed']], "Machine Speed - Idle Periods", '#ff0000')
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.line_chart(idle_data.set_index('timestamp')[['speed']], color='#ff0000', use_container_width=True)
    else:
        st.warning("No idle periods found in the selected time range.")

elif graph_view == "Low Speed Events":
    st.subheader("Speed Over Time - Low Speed Events Highlighted")
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=filtered_df['timestamp'], y=filtered_df['speed'], mode='lines', name='Speed', line=dict(color='#1f77b4')))
    low_speed_data = filtered_df[filtered_df['is_low_speed']]
    if not low_speed_data.empty:
        fig.add_trace(go.Scatter(x=low_speed_data['timestamp'], y=low_speed_data['speed'], mode='markers', name='Low Speed', marker=dict(color='orange', size=8)))
    fig.update_layout(title=f"Machine Speed with Low Speed Events (1 < Speed < {mc_min_speed_req} RPM) Highlighted", xaxis_title="Time", yaxis_title="Speed (RPM)")
    st.plotly_chart(fig, use_container_width=True)

elif graph_view == "Production vs Speed":
    st.subheader("Production and Speed Correlation")
    filtered_df['production_rate'] = filtered_df['quantity'].diff().fillna(0)
    chart_data = filtered_df.set_index('timestamp')[['speed', 'production_rate']]
    st.line_chart(chart_data)
    if len(filtered_df) > 1:
        correlation = filtered_df['speed'].corr(filtered_df['production_rate'])
        st.write(f"**Speed-Production Correlation:** {correlation:.3f}")
        
elif graph_view == "Idle Events Timeline":
    st.subheader("Idle Events Timeline Visualization")
    if not filtered_idle_events.empty:
        gantt_data = [dict(Task=f"Idle Group {idx}", Start=row['start'], Finish=row['end'], Resource=f"{row['duration']:.1f} min") for idx, row in filtered_idle_events.iterrows()]
        if gantt_data:
            fig = ff.create_gantt(gantt_data, index_col='Resource', title='Idle Events Timeline', show_colorbar=True, bar_width=0.4, showgrid_x=True, showgrid_y=True)
            st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("No idle events to visualize in the selected time range.")
        
elif graph_view == "Hourly Performance Pattern":
    st.subheader("Hourly Performance Pattern")
    if chart_type == "Bar Chart":
        hourly_chart_data = hourly_stats[['Avg_Speed', 'Uptime_Percent']].copy()
        fig = create_multi_bar_chart(hourly_chart_data, ['Avg_Speed', 'Uptime_Percent'], "Hourly Performance - Average Speed and Uptime")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.line_chart(hourly_stats[['Avg_Speed', 'Uptime_Percent']], use_container_width=True)


# --- DEEPSEEK AI SECTION ---
st.header("DeepSeek AI Analysis")

def ask_deepseek(prompt, temperature=0.1):
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    enhanced_system_prompt = "You are a senior IoT data analyst..."
    data = {"model": MODEL, "messages": [{"role": "system", "content": enhanced_system_prompt}, {"role": "user", "content": prompt}], "temperature": temperature, "max_tokens": 2000, "top_p": 0.9}
    try:
        response = requests.post(DEEPSEEK_URL, headers=headers, json=data, timeout=300)
        return response.json()["choices"][0]["message"]["content"] if response.status_code == 200 else f"DeepSeek API Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Error connecting to DeepSeek API: {str(e)}"

st.subheader("AI-Powered Pattern Analysis (DeepSeek) - Dataframe Driven")
analysis_option = st.radio("Select analysis scope:", ["Full Dataset Analysis", "Filtered Data Analysis", "Time Pattern Analysis"])
if st.button("🤖 Analyze Data with DeepSeek AI", type="primary"):
    st.write("...AI analysis will be displayed here...")

st.subheader("Ask DeepSeek AI About Your Data")
user_query = st.text_input("💬 Ask a question about your machine data")
if user_query:
    st.write("...AI answer will be displayed here...")

st.subheader("Advanced AI Analysis Options")
col1, col2 = st.columns(2)
with col1:
    if st.button("🔧 Maintenance Schedule Recommendations"):
        st.write("...Maintenance recommendations will be displayed here...")
with col2:
    if st.button("📈 Production Optimization Analysis"):
        st.write("...Production optimization analysis will be displayed here...")


# --- DOWNLOAD BUTTON SECTION ---
st.header("Generate & Download Report")

def create_comprehensive_report_with_graphs(df, filtered_df, min_speed_req, current_idle_time, kpi_data, selected_file):
    doc = Document()
    doc.add_heading('IoT Machine Analytics Report', 0)
    # This is a placeholder as the full implementation was not in the original file
    doc.add_paragraph(f"Report for: {selected_file}")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_heading('Key Performance Indicators', level=1)
    doc.add_paragraph(f"Uptime: {kpi_data.get('uptime', 0):.1f}%")
    doc.add_paragraph(f"Average Speed (when running): {kpi_data.get('avg_speed', 0):.1f} RPM")
    doc.add_paragraph(f"Total Production: {kpi_data.get('total_production', 0):.1f}")
    doc.add_paragraph(f"Long Idle Events (> {current_idle_time} min): {kpi_data.get('idle_events', 0)}")
    # Add more sections for tables and graphs as needed
    return doc

if st.button("Download Report (DOCX)"):
    if 'df' in locals() and 'filtered_df' in locals():
        kpi_data_for_report = {'uptime': uptime_percent, 'avg_speed': avg_speed, 'total_production': total_production, 'idle_events': len(long_idles)}
        doc = create_comprehensive_report_with_graphs(df, filtered_df, mc_min_speed_req, idle_time, kpi_data_for_report, selected_file)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        st.download_button(
            label="Click to Download Report",
            data=bio.getvalue(),
            file_name=f"Machine_Analytics_Report_{selected_file.replace('.csv', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Please select and analyze data first to generate a report.")